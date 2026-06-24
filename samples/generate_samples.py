from docx import Document
from docx.shared import Pt, Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
import os

# Register a Chinese font. Try common Windows Chinese fonts.
font_paths = [
    r'C:\Windows\Fonts\msyh.ttc',
    r'C:\Windows\Fonts\simsun.ttc',
    r'C:\Windows\Fonts\simhei.ttf',
]
chinese_font = None
for fp in font_paths:
    if os.path.exists(fp):
        try:
            name = 'ChineseFont'
            pdfmetrics.registerFont(TTFont(name, fp))
            chinese_font = name
            break
        except Exception:
            continue

if chinese_font is None:
    raise RuntimeError('未找到可用的中文字体，请安装微软雅黑或宋体。')

samples = {
    '房屋租赁合同': '''房屋租赁合同

出租方（甲方）：王先生
承租方（乙方）：李女士

根据《中华人民共和国民法典》及相关法律法规，甲乙双方本着平等、自愿、公平的原则，就房屋租赁事宜达成如下协议：

第一条 房屋基本情况
1.1 甲方将位于北京市朝阳区阳光小区 3 号楼 2 单元 1502 室的房屋出租给乙方使用，建筑面积约为 89 平方米。
1.2 该房屋用途为居住，未经甲方书面同意，乙方不得擅自改变房屋用途。

第二条 租赁期限
2.1 租赁期限自 2026 年 7 月 1 日起至 2027 年 6 月 30 日止，共计 12 个月。

第三条 租金及支付方式
3.1 该房屋月租金为人民币捌仟元整（¥8,000）。
3.2 乙方应于每月 5 日前将当月租金支付至甲方指定账户。
3.3 逾期支付租金的，每逾期一日，乙方应按月租金的 10% 向甲方支付违约金。

第四条 押金
4.1 乙方应于签订本合同时向甲方支付押金人民币壹万陆仟元整（¥16,000）。
4.2 租赁期满或合同解除后，甲方应在乙方交还房屋并结清相关费用后七日内将押金无息退还乙方。如房屋及设施有损坏，甲方有权从押金中扣除相应维修费用。

第五条 房屋维护
5.1 租赁期间，房屋的日常维护由乙方负责；房屋主体结构及重大维修由甲方负责。
5.2 因乙方使用不当造成房屋及设施损坏的，乙方应负责修复或赔偿。

第六条 合同解除
6.1 乙方有下列情形之一的，甲方有权单方解除合同：
（1）未按约定支付租金达 3 日以上的；
（2）擅自改变房屋用途的；
（3）擅自转租、分租的。
6.2 租赁期内，任何一方提前解除合同的，应提前 3 日通知对方，并支付相当于 1 个月租金的违约金。

第七条 争议解决
7.1 本合同履行过程中发生争议的，双方应协商解决；协商不成的，任何一方均可向甲方所在地人民法院提起诉讼。

第八条 其他
8.1 本合同一式两份，甲乙双方各执一份，自双方签字之日起生效。

甲方（签字）：____________    乙方（签字）：____________
日期：____年____月____日      日期：____年____月____日
''',
    '劳动合同': '''劳动合同

甲方（用人单位）：星辰科技有限公司
乙方（劳动者）：张明

根据《中华人民共和国劳动法》《中华人民共和国劳动合同法》及相关法律法规，甲乙双方经平等协商，自愿签订本合同。

第一条 合同期限
1.1 本合同为固定期限劳动合同，自 2026 年 7 月 1 日起至 2029 年 6 月 30 日止，共计 3 年。
1.2 试用期为 6 个月，自 2026 年 7 月 1 日起至 2026 年 12 月 31 日止。

第二条 工作内容与地点
2.1 乙方同意根据甲方工作需要，担任软件工程师岗位。
2.2 乙方的工作地点为甲方指定办公场所，甲方可根据经营需要随时调整乙方工作地点，乙方应无条件服从。

第三条 工作时间与休息休假
3.1 乙方实行标准工时制，每日工作 8 小时，每周工作 5 天。
3.2 甲方因生产经营需要安排乙方加班的，乙方应予以配合，甲方按法律规定支付加班费或安排调休。

第四条 劳动报酬
4.1 乙方试用期月工资为人民币壹万贰仟元整（¥12,000），转正后月工资为人民币壹万伍仟元整（¥15,000）。
4.2 甲方应于每月 10 日前以货币形式支付乙方上月工资。
4.3 乙方离职时，甲方有权暂扣乙方最后一个月工资，待乙方完成工作交接并无遗留问题后支付。

第五条 社会保险与福利
5.1 甲方按国家和地方规定为乙方缴纳社会保险和住房公积金，乙方个人缴纳部分由甲方从其工资中代扣代缴。

第六条 保密与竞业限制
6.1 乙方应对在工作中知悉的甲方商业秘密予以保密，保密义务在离职后永久有效。
6.2 乙方离职后 3 年内不得在与甲方有竞争关系的单位任职，甲方无需向乙方支付竞业限制补偿金。

第七条 合同解除与终止
7.1 乙方有下列情形之一的，甲方可以解除劳动合同：
（1）严重违反甲方规章制度的；
（2）严重失职，营私舞弊，给甲方造成重大损害的；
（3）被依法追究刑事责任的。
7.2 乙方提前解除劳动合同的，应提前 30 日以书面形式通知甲方。

第八条 争议解决
8.1 因履行本合同发生争议的，双方应协商解决；协商不成的，可向劳动争议仲裁委员会申请仲裁。

第九条 其他
9.1 本合同未尽事宜，按照国家有关法律、法规和甲方依法制定的规章制度执行。
9.2 本合同一式两份，甲乙双方各执一份，自双方签字盖章之日起生效。

甲方（盖章）：____________    乙方（签字）：____________
日期：____年____月____日      日期：____年____月____日
''',
    '货物买卖合同': '''货物买卖合同

买方（甲方）：东方贸易集团有限公司
卖方（乙方）：华南制造有限公司

根据《中华人民共和国民法典》及相关法律法规，甲乙双方本着平等互利原则，就甲方向乙方购买货物事宜达成如下协议：

第一条 货物名称、规格、数量及价格
1.1 甲方向乙方采购办公家具一批，具体规格、数量及单价详见附件《货物清单》。
1.2 本合同总金额为人民币贰佰万元整（¥2,000,000），含税。

第二条 质量标准
2.1 货物应符合国家相关质量标准及双方约定的技术要求。
2.2 乙方应随货提供产品合格证、质量检测报告等相关文件。

第三条 交货时间与地点
3.1 乙方应于合同签订后 90 日内将货物运送至甲方指定地点。
3.2 运输费用由乙方承担，运输途中货物毁损、灭失的风险由乙方承担。

第四条 验收标准及异议期限
4.1 甲方应在收到货物后 3 日内完成验收，逾期未提出书面异议的，视为验收合格。
4.2 验收不合格的，甲方有权要求乙方在 15 日内更换或修复。

第五条 付款方式
5.1 甲方应于合同签订后 5 个工作日内支付合同总价款的 30% 作为预付款。
5.2 货物验收合格后，甲方应在 180 日内支付剩余 70% 货款。
5.3 乙方应在甲方付款前向甲方开具等额增值税专用发票。

第六条 违约责任
6.1 乙方逾期交货的，每逾期一日，应按合同总金额的 1% 向甲方支付违约金；逾期超过 30 日的，甲方有权解除合同。
6.2 甲方逾期付款的，每逾期一日，应按未付款项的 0.5% 向乙方支付违约金。
6.3 任何一方违反本合同约定的，应向守约方支付合同总金额 40% 的违约金。

第七条 知识产权
7.1 乙方保证所供货物不侵犯任何第三方知识产权，否则由乙方承担全部法律责任。

第八条 争议解决
8.1 因本合同引起的或与本合同有关的任何争议，双方应友好协商解决；协商不成的，提交甲方所在地仲裁委员会仲裁解决。

第九条 其他
9.1 本合同自双方签字盖章之日起生效，有效期至合同履行完毕之日止。
9.2 本合同一式两份，甲乙双方各执一份，具有同等法律效力。

甲方（盖章）：____________    乙方（盖章）：____________
日期：____年____月____日      日期：____年____月____日
'''
}


def create_docx(title, text, output_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    # For Chinese font setting in python-docx, also set eastAsia
    style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/drawingml/2006/main}eastAsia', '宋体')

    heading = doc.add_heading(title, level=1)
    heading.alignment = 1  # center

    for line in text.split('\n'):
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)

    doc.save(output_path)


def create_pdf(title, text, output_path):
    c = canvas.Canvas(output_path, pagesize=A4)
    width, height = A4
    margin = 50
    max_width = width - 2 * margin
    line_height = 20
    y = height - margin

    # Title
    c.setFont(chinese_font, 18)
    title_width = c.stringWidth(title, chinese_font, 18)
    c.drawString((width - title_width) / 2, y, title)
    y -= line_height * 2

    c.setFont(chinese_font, 11)
    for raw_line in text.split('\n')[1:]:
        line = raw_line.rstrip()
        if not line:
            y -= line_height * 0.6
            continue

        # Wrap long lines manually
        chars = []
        current_width = 0
        for char in line:
            char_width = c.stringWidth(char, chinese_font, 11)
            if current_width + char_width > max_width and chars:
                c.drawString(margin, y, ''.join(chars))
                y -= line_height
                chars = [char]
                current_width = char_width
            else:
                chars.append(char)
                current_width += char_width
        if chars:
            c.drawString(margin, y, ''.join(chars))
            y -= line_height

        if y < margin + line_height:
            c.showPage()
            c.setFont(chinese_font, 11)
            y = height - margin

    c.save()


def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    for title, text in samples.items():
        docx_path = os.path.join(base_dir, f'{title}.docx')
        pdf_path = os.path.join(base_dir, f'{title}.pdf')
        create_docx(title, text, docx_path)
        create_pdf(title, text, pdf_path)
        print(f'已生成：{docx_path}')
        print(f'已生成：{pdf_path}')


if __name__ == '__main__':
    main()
